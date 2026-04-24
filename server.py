import base64
import json
import os
import re
import subprocess
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote

import keyring
import msal
import requests
from docx import Document
from mcp.server.fastmcp import FastMCP

# ── Constants ──────────────────────────────────────────────────────────────────

_ADO_SCOPE = "499b84ac-1321-427f-aa17-267ca6975798/.default"
_MSAL_SCOPES = ["499b84ac-1321-427f-aa17-267ca6975798/user_impersonation"]

# VS Code public client ID — pre-consented in most Azure AD tenants, supports
# device-code flows to Azure DevOps (same identity stack as VS Code extensions).
_MSAL_CLIENT_ID = "aebc6443-996d-45c2-90f0-388ff96faa56"
_MSAL_AUTHORITY = "https://login.microsoftonline.com/organizations"

# Persistent state (token cache, in-progress device flow, config)
_STATE_DIR = Path.home() / ".mcp-docx-server"
_CACHE_PATH = _STATE_DIR / "token_cache.json"
_FLOW_PATH = _STATE_DIR / "device_flow.json"
_CONFIG_PATH = _STATE_DIR / "config.json"

# PAT fallback
_KEYRING_SERVICE = "mcp-docx-server"
_KEYRING_USERNAME = "ado-pat"

mcp = FastMCP("mcp-docx-server")


# ── MSAL helpers ───────────────────────────────────────────────────────────────

def _load_config() -> dict:
    if _CONFIG_PATH.exists():
        try:
            return json.loads(_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_config(cfg: dict) -> None:
    _STATE_DIR.mkdir(parents=True, exist_ok=True)
    _CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


def _get_authority() -> str:
    tenant = _load_config().get("tenant")
    if tenant:
        return f"https://login.microsoftonline.com/{tenant}"
    return _MSAL_AUTHORITY


def _load_cache() -> msal.SerializableTokenCache:
    cache = msal.SerializableTokenCache()
    if _CACHE_PATH.exists():
        try:
            cache.deserialize(_CACHE_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return cache


def _save_cache(cache: msal.SerializableTokenCache) -> None:
    if cache.has_state_changed:
        _STATE_DIR.mkdir(parents=True, exist_ok=True)
        _CACHE_PATH.write_text(cache.serialize(), encoding="utf-8")


def _get_ado_token_silent() -> str | None:
    """Return any cached ADO Bearer token (used when no org context is available)."""
    cache = _load_cache()
    app = msal.PublicClientApplication(
        _MSAL_CLIENT_ID, authority=_get_authority(), token_cache=cache
    )
    accounts = app.get_accounts()
    if not accounts:
        return None
    result = app.acquire_token_silent(scopes=_MSAL_SCOPES, account=accounts[0])
    if result and "access_token" in result:
        _save_cache(cache)
        return result["access_token"]
    return None


def _discover_tenant(org_name: str) -> str | None:
    """
    Discover the Azure AD tenant for an ADO organisation.
    Uses allow_redirects=False so the 302 response headers are visible;
    reads X-VSS-ResourceTenant (most reliable) or parses WWW-Authenticate.
    """
    try:
        resp = requests.get(
            f"https://dev.azure.com/{org_name}",
            headers={"Authorization": "Bearer invalid"},
            allow_redirects=False,
            timeout=10,
        )
        # Most reliable: dedicated tenant header
        tenant = resp.headers.get("X-VSS-ResourceTenant", "").strip()
        if tenant:
            return tenant
        # Fallback: parse WWW-Authenticate
        auth_header = resp.headers.get("WWW-Authenticate", "")
        m = re.search(
            r"authorization_uri=https://login\.(?:windows|microsoftonline)\.net/([^,\s]+)",
            auth_header,
        )
        if m:
            return m.group(1)
    except Exception:
        pass
    return None


def _get_token_for_org(org_name: str) -> str | tuple[None, str]:
    """
    Return a cached ADO Bearer token for the tenant that owns *org_name*.
    Discovers the tenant automatically; strictly requires a cached account
    from that exact tenant.
    Returns the token string on success, or (None, error_message) if not
    authenticated for that org's tenant.
    """
    tenant = _discover_tenant(org_name)
    if not tenant:
        # Can't discover tenant — fall back to any cached token
        cache = _load_cache()
        app = msal.PublicClientApplication(
            _MSAL_CLIENT_ID, authority=_get_authority(), token_cache=cache
        )
        accounts = app.get_accounts()
        if not accounts:
            return None, (
                f"Not signed in. Call login_ado_start(org='{org_name}') to authenticate."
            )
        result = app.acquire_token_silent(scopes=_MSAL_SCOPES, account=accounts[0])
        if result and "access_token" in result:
            _save_cache(cache)
            return result["access_token"], ""
        return None, f"Token refresh failed. Call login_ado_start(org='{org_name}') to re-authenticate."

    authority = f"https://login.microsoftonline.com/{tenant}"
    cache = _load_cache()
    app = msal.PublicClientApplication(
        _MSAL_CLIENT_ID, authority=authority, token_cache=cache
    )
    # Strictly only use accounts whose home tenant matches the discovered tenant
    matching = [
        acc for acc in app.get_accounts()
        if tenant.lower() in acc.get("home_account_id", "").lower()
    ]
    if not matching:
        return None, (
            f"Not signed in for organisation '{org_name}' (tenant {tenant}). "
            f"Call login_ado_start(org='{org_name}') to authenticate with the right account."
        )
    result = app.acquire_token_silent(scopes=_MSAL_SCOPES, account=matching[0])
    if result and "access_token" in result:
        _save_cache(cache)
        return result["access_token"], ""
    return None, (
        f"Session expired for organisation '{org_name}'. "
        f"Call login_ado_start(org='{org_name}') to re-authenticate."
    )


def _parse_workitem_url(url: str) -> tuple[str, str, int] | None:
    """
    Parse an ADO work item URL into (org, project, work_item_id).
    Supports:
      https://dev.azure.com/{org}/{project}/_workitems/edit/{id}
    Returns None if the URL doesn't match.
    """
    m = re.match(
        r"https://dev\.azure\.com/([^/]+)/([^/]+)/_workitems/(?:edit|view)/(\d+)",
        url.rstrip("/"),
    )
    if not m:
        return None
    return unquote(m.group(1)), unquote(m.group(2)), int(m.group(3))


# ── Fetch helper ───────────────────────────────────────────────────────────────

def _get_az_cli_token() -> str | None:
    """
    Get an ADO Bearer token via Azure CLI (az account get-access-token).
    Works when the user is already signed in with `az login` — same auth
    as used by @azure-devops/mcp.
    """
    try:
        result = subprocess.run(
            [
                "az", "account", "get-access-token",
                "--resource", "499b84ac-1321-427f-aa17-267ca6975798",
                "--query", "accessToken",
                "-o", "tsv",
            ],
            capture_output=True,
            text=True,
            timeout=15,
        )
        if result.returncode == 0:
            token = result.stdout.strip()
            if token:
                return token
    except Exception:
        pass
    return None


def _get_gcm_token(org_name: str) -> str | None:
    """
    Ask Git Credential Manager for a token for the given ADO org.
    GCM already holds working credentials (same as used by git clone/push)
    and handles all auth complexities including Conditional Access policies.
    Returns a Basic-auth compatible password string, or None if unavailable.
    """
    try:
        input_str = f"protocol=https\nhost=dev.azure.com\nusername={org_name}\n\n"
        result = subprocess.run(
            ["git", "credential", "fill"],
            input=input_str,
            capture_output=True,
            text=True,
            timeout=10,
        )
        if result.returncode != 0:
            return None
        creds = {}
        for line in result.stdout.splitlines():
            if "=" in line:
                k, _, v = line.partition("=")
                creds[k.strip()] = v.strip()
        return creds.get("password")
    except Exception:
        return None


def _fetch_ado_url(url: str, org_name: str = "") -> bytes:
    """
    Fetch a URL from Azure DevOps. Auth priority:
    1. MSAL cached token for the org's tenant (strictly matched)
    2. Azure CLI token (az account get-access-token) — same auth as @azure-devops/mcp
    3. Git Credential Manager token (git credential fill) — works without any setup
    4. PAT from Windows Credential Manager (store_ado_pat) or ADO_PAT env var
    """
    if org_name:
        token, err = _get_token_for_org(org_name)
    else:
        token, err = _get_ado_token_silent(), ""

    if token:
        resp = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
        if resp.ok:
            return resp.content
        # 401/403 — token may not cover this org, fall through to other methods

    # MSAL failed — try Azure CLI (zero-setup if `az login` already done)
    az_token = _get_az_cli_token()
    if az_token:
        resp = requests.get(url, headers={"Authorization": f"Bearer {az_token}"}, timeout=30)
        if resp.ok:
            return resp.content

    # Azure CLI failed — try Git Credential Manager (zero-setup, uses existing git auth)
    if org_name:
        gcm_password = _get_gcm_token(org_name)
        if gcm_password:
            resp = requests.get(url, auth=("", gcm_password), timeout=30)
            if resp.ok:
                return resp.content

    # GCM failed — try PAT fallback
    pat = keyring.get_password(_KEYRING_SERVICE, _KEYRING_USERNAME) or os.environ.get("ADO_PAT", "")
    if not pat:
        raise RuntimeError(err or (
            "Not authenticated with Azure DevOps. "
            "Either run 'git clone https://dev.azure.com/{org}' once to set up Git Credential Manager, "
            "or call login_ado_start() to sign in."
        ))
    resp = requests.get(url, auth=("", pat), timeout=30)
    resp.raise_for_status()
    return resp.content


# ── Tools ──────────────────────────────────────────────────────────────────────

@mcp.tool()
def login_ado_start(org: str = "", tenant: str = "") -> str:
    """
    Start Azure DevOps sign-in using a browser-based device code flow.
    No PAT required — same auth approach used by @azure-devops/mcp.

    Recommended: pass the ADO organisation name so the correct tenant is
    auto-discovered and the sign-in targets exactly the right account:
        login_ado_start(org="Nobilis-Group")

    If you have multiple ADO organisations in different tenants, call this
    once per organisation. Each session is cached separately.

    org:    ADO organisation name (e.g. "Nobilis-Group"). Tenant is
            auto-discovered from this — no need to look it up manually.
    tenant: explicit tenant ID override (only needed if org discovery fails).

    After calling this, open the URL in a browser, enter the code, sign in,
    then call login_ado_complete().
    """
    # Resolve authority: org discovery > explicit tenant > saved config > default
    if org:
        discovered = _discover_tenant(org)
        if discovered:
            authority = f"https://login.microsoftonline.com/{discovered}"
        elif tenant:
            authority = f"https://login.microsoftonline.com/{tenant}"
        else:
            authority = _get_authority()
    elif tenant:
        cfg = _load_config()
        cfg["tenant"] = tenant
        _save_config(cfg)
        authority = f"https://login.microsoftonline.com/{tenant}"
    else:
        authority = _get_authority()

    cache = _load_cache()
    app = msal.PublicClientApplication(
        _MSAL_CLIENT_ID, authority=authority, token_cache=cache
    )
    flow = app.initiate_device_flow(scopes=_MSAL_SCOPES)
    if "error" in flow:
        return f"Error starting sign-in: {flow.get('error_description', flow['error'])}"
    _STATE_DIR.mkdir(parents=True, exist_ok=True)
    # Store the authority used so login_ado_complete can reconstruct the same app
    flow["_authority"] = authority
    _FLOW_PATH.write_text(json.dumps(flow), encoding="utf-8")
    return f"{flow['message']}\n\n(signing in for: {authority})"


@mcp.tool()
def login_ado_complete() -> str:
    """
    Complete the Azure DevOps sign-in after entering the device code in the browser.
    Call login_ado_start() first to get the code.
    If you haven't finished the browser step yet, complete it first and call this again.
    """
    if not _FLOW_PATH.exists():
        return "No pending sign-in found. Call login_ado_start() first."

    flow = json.loads(_FLOW_PATH.read_text(encoding="utf-8"))
    authority = flow.pop("_authority", _get_authority())
    cache = _load_cache()
    app = msal.PublicClientApplication(
        _MSAL_CLIENT_ID, authority=authority, token_cache=cache
    )

    # Poll the token endpoint once (exit_condition=lambda _: True stops after the
    # first attempt). Takes ~5 s (one sleep interval). If the browser step isn't
    # done yet the result will be "authorization_pending" and the user can retry.
    result = app.acquire_token_by_device_flow(
        flow,
        exit_condition=lambda elapsed: True,
    )

    if result and "access_token" in result:
        _save_cache(cache)
        _FLOW_PATH.unlink(missing_ok=True)
        return (
            "Azure DevOps sign-in successful. "
            "Session cached — read_docx_ado now works without any PAT."
        )

    error = (result or {}).get("error", "unknown")
    if error == "authorization_pending":
        return (
            "Still waiting — complete the browser sign-in first, "
            "then call login_ado_complete() again."
        )
    return f"Sign-in failed: {(result or {}).get('error_description', error)}"


@mcp.tool()
def store_ado_pat(pat: str) -> str:
    """
    Alternative to device-code login: save an Azure DevOps Personal Access Token
    in Windows Credential Manager (encrypted by the OS, never written to a file).
    The PAT needs at least Work Items - Read scope.
    Call this once — read_docx_ado will use it automatically as a fallback.
    """
    if not pat:
        return "Error: PAT must not be empty."
    keyring.set_password(_KEYRING_SERVICE, _KEYRING_USERNAME, pat)
    return "PAT saved to Windows Credential Manager."


@mcp.tool()
def read_docx_from_workitem(workitem_url: str) -> str:
    """
    Given an Azure DevOps work item URL, find all .docx attachments on that
    work item and return their full text content.

    Example URL: https://dev.azure.com/Nobilis-Group/Business%20Central%20365/_workitems/edit/1858/

    Authentication is automatic — uses the session from login_ado_start/complete.
    The correct tenant is detected automatically, so this works across multiple
    ADO organisations without any extra configuration.
    """
    parsed = _parse_workitem_url(workitem_url)
    if not parsed:
        return (
            "Error: Could not parse work item URL. "
            "Expected: https://dev.azure.com/{org}/{project}/_workitems/edit/{id}"
        )
    org, _project, wi_id = parsed

    # Fetch work item with all relations (attachments live there)
    api_url = (
        f"https://dev.azure.com/{org}/_apis/wit/workitems/{wi_id}"
        f"?$expand=relations&api-version=7.1"
    )
    try:
        raw = _fetch_ado_url(api_url, org_name=org)
    except RuntimeError as exc:
        return f"Error: {exc}"
    except requests.RequestException as exc:
        return f"Error: Could not fetch work item: {exc}"

    try:
        wi = json.loads(raw)
    except Exception as exc:
        return f"Error: Could not parse work item response: {exc}"

    relations = wi.get("relations") or []
    attachments = [
        r for r in relations
        if r.get("rel") == "AttachedFile"
        and r.get("attributes", {}).get("name", "").lower().endswith(".docx")
    ]

    if not attachments:
        names = [
            r.get("attributes", {}).get("name", "?")
            for r in relations
            if r.get("rel") == "AttachedFile"
        ]
        if names:
            return (
                f"No .docx attachments found on work item {wi_id}. "
                f"Attachments present: {', '.join(names)}"
            )
        return f"No attachments found on work item {wi_id}."

    results = []
    for att in attachments:
        name = att.get("attributes", {}).get("name", "attachment.docx")
        att_url = att["url"]
        try:
            docx_bytes = _fetch_ado_url(att_url, org_name=org)
            document = Document(BytesIO(docx_bytes))
            text = "\n".join(p.text for p in document.paragraphs)
            results.append(f"=== {name} ===\n{text}")
        except Exception as exc:
            results.append(f"=== {name} ===\nError reading file: {exc}")

    return "\n\n".join(results)


@mcp.tool()
def read_docx_local(file_path: str) -> str:
    """Read text content from a local .docx file."""
    try:
        document = Document(Path(file_path))
        return "\n".join(p.text for p in document.paragraphs)
    except FileNotFoundError:
        return f"Error: File not found: {file_path}"
    except Exception as exc:
        return f"Error: Could not read DOCX file '{file_path}': {exc}"


@mcp.tool()
def read_docx_ado(attachment_url: str) -> str:
    """
    Download a .docx attachment from Azure DevOps and return its full text.
    Uses the session from login_ado_start/complete — no PAT in config or prompts.
    Falls back to a stored PAT (store_ado_pat) if no session is active.
    """
    try:
        raw = _fetch_ado_url(attachment_url)
    except RuntimeError as exc:
        return f"Error: {exc}"
    except requests.RequestException as exc:
        return f"Error: Could not download attachment: {exc}"
    try:
        document = Document(BytesIO(raw))
        return "\n".join(p.text for p in document.paragraphs)
    except Exception as exc:
        return f"Error: Could not parse DOCX content: {exc}"


@mcp.tool()
def read_docx_bytes(base64_content: str) -> str:
    """
    Parse a .docx file from base64-encoded binary content and return its text.
    Use this when another tool has already fetched the file as a base64 string.
    """
    try:
        raw = base64.b64decode(base64_content)
    except Exception as exc:
        return f"Error: Could not decode base64 content: {exc}"
    try:
        document = Document(BytesIO(raw))
        return "\n".join(p.text for p in document.paragraphs)
    except Exception as exc:
        return f"Error: Could not parse DOCX content: {exc}"


if __name__ == "__main__":
    mcp.run(transport="stdio")
