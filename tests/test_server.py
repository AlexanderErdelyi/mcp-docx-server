from io import BytesIO
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import Mock, patch

from docx import Document
import requests

from server import read_docx_ado, read_docx_local


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestServerTools(unittest.TestCase):
    def test_read_docx_local_reads_heading_and_paragraph(self) -> None:
        with TemporaryDirectory() as tmpdir:
            file_path = f"{tmpdir}/sample.docx"
            document = Document()
            document.add_heading("Heading 1", level=1)
            document.add_paragraph("First paragraph.")
            document.save(file_path)

            result = read_docx_local(file_path)

        self.assertEqual(result, "Heading 1\nFirst paragraph.")

    def test_read_docx_local_file_not_found(self) -> None:
        result = read_docx_local("/does/not/exist.docx")
        self.assertTrue(result.startswith("Error: File not found:"))

    @patch("server.requests.get")
    def test_read_docx_ado_downloads_and_parses(self, mock_get: Mock) -> None:
        response = Mock()
        response.content = _docx_bytes("A", "B")
        response.raise_for_status.return_value = None
        mock_get.return_value = response

        result = read_docx_ado("https://dev.azure.com/x/y", "pat-token")

        mock_get.assert_called_once_with(
            "https://dev.azure.com/x/y", auth=("", "pat-token"), timeout=30
        )
        self.assertEqual(result, "A\nB")

    @patch("server.requests.get")
    def test_read_docx_ado_handles_download_error(self, mock_get: Mock) -> None:
        mock_get.side_effect = requests.RequestException("boom")

        result = read_docx_ado("https://dev.azure.com/x/y", "pat-token")

        self.assertTrue(result.startswith("Error: Could not download attachment:"))

    @patch("server.requests.get")
    def test_read_docx_ado_requires_pat(self, mock_get: Mock) -> None:
        result = read_docx_ado("https://dev.azure.com/x/y", "")
        self.assertEqual(result, "Error: Azure DevOps PAT is required.")
        mock_get.assert_not_called()


if __name__ == "__main__":
    unittest.main()
